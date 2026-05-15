"""
文档加载器 - 支持 PDF / Word / Markdown / TXT
返回 LangChain Document 对象（page_content + metadata）。
"""

import os
from pathlib import Path
from typing import List, Optional

from langchain_core.documents import Document


def _load_pdf(file_path: Path) -> List[Document]:
    import fitz  # pymupdf

    docs = []
    try:
        doc = fitz.open(str(file_path))
        for page_num in range(len(doc)):
            text = doc[page_num].get_text()
            if text.strip():
                docs.append(Document(
                    page_content=text,
                    metadata={"source": str(file_path), "page": page_num + 1, "type": "pdf"},
                ))
        doc.close()
    except Exception as e:
        print(f"Warning: failed to load PDF {file_path}: {e}")
    return docs


def _load_docx(file_path: Path) -> List[Document]:
    from docx import Document as DocxDocument

    docs = []
    try:
        doc = DocxDocument(str(file_path))
        text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        if text.strip():
            docs.append(Document(
                page_content=text,
                metadata={"source": str(file_path), "type": "docx"},
            ))
    except Exception as e:
        print(f"Warning: failed to load DOCX {file_path}: {e}")
    return docs


def _load_markdown(file_path: Path) -> List[Document]:
    docs = []
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read()
        if text.strip():
            docs.append(Document(
                page_content=text,
                metadata={"source": str(file_path), "type": "markdown"},
            ))
    except Exception as e:
        print(f"Warning: failed to load Markdown {file_path}: {e}")
    return docs


def _load_txt(file_path: Path) -> List[Document]:
    docs = []
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read()
        if text.strip():
            docs.append(Document(
                page_content=text,
                metadata={"source": str(file_path), "type": "txt"},
            ))
    except Exception as e:
        print(f"Warning: failed to load TXT {file_path}: {e}")
    return docs


_LOADERS = {
    ".pdf": _load_pdf,
    ".docx": _load_docx,
    ".md": _load_markdown,
    ".txt": _load_txt,
    ".markdown": _load_markdown,
}


def load_documents(directory: Optional[str] = None) -> List[Document]:
    """加载指定目录下所有支持的文档，返回 LangChain Document 列表。"""
    if directory is None:
        from app.config import settings
        directory = settings.knowledge_base_dir

    path = Path(directory)
    if not path.exists():
        return []

    documents: List[Document] = []
    for file_path in sorted(path.rglob("*")):
        if not file_path.is_file():
            continue
        ext = file_path.suffix.lower()
        loader = _LOADERS.get(ext)
        if loader:
            documents.extend(loader(file_path))

    return documents
